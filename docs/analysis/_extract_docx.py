"""Extract text from all Desktop .docx files into individual .txt files for analysis."""
import os, sys
from pathlib import Path
from docx import Document

desktop = Path(os.environ["USERPROFILE"]) / "Desktop"
out_dir = Path(__file__).parent / "_docx_extracts"
out_dir.mkdir(exist_ok=True)

for f in sorted(desktop.glob("*.docx")):
    if f.stat().st_size == 0:
        print(f"SKIP (empty): {f.name}")
        continue
    try:
        doc = Document(str(f))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Also extract tables
        for i, tbl in enumerate(doc.tables):
            text += f"\n\n--- TABLE {i+1} ---\n"
            for row in tbl.rows:
                text += " | ".join(c.text.strip() for c in row.cells) + "\n"
        out_path = out_dir / f"{f.stem}.txt"
        out_path.write_text(text, encoding="utf-8")
        lines = len(text.splitlines())
        print(f"OK: {f.name} -> {out_path.name} ({lines} lines, {len(text)} chars)")
    except Exception as e:
        print(f"ERROR: {f.name}: {e}")

print(f"\nAll extracts in: {out_dir}")
